import pytest
import os
from io import BytesIO
from docx import Document
from app.schemas.menu import MenuRequest, MenuData, SubCategoryItem
from app.services.docx_generator import generate_menu_docx

def test_generate_menu_docx():
    # Setup test data
    payload = {
        "all_meals": [
            {
                "categoria": "Desayuno",
                "fecha": "2023-11-01",
                "descripcion": "Desayuno completo",
                "items": [
                    {"subcat": "Frutas", "menu": "Manzana, Pera"},
                    {"subcat": "Bebidas", "menu": "Café, Jugo"}
                ]
            }
        ]
    }
    menu_request = MenuRequest(**payload)
    
    # Generate document
    docx_stream = generate_menu_docx(menu_request)
    
    # Assert return type
    assert isinstance(docx_stream, BytesIO)
    
    # Load document from stream to verify contents
    docx_stream.seek(0)
    doc = Document(docx_stream)
    
    # Very basic verification: The text should be present somewhere in the document's paragraphs or tables
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    
    # Also check tables if the generator uses them
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text.append(cell.text)
                
    content_str = " ".join(full_text)
    
    # Validate content elements exist
    assert "Desayuno" in content_str
    assert "2023-11-01" in content_str
    assert "Desayuno completo" in content_str
    assert "Frutas" in content_str
    assert "Manzana, Pera" in content_str
