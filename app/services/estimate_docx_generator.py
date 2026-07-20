import os
import logging
import re
from datetime import datetime
from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from app.schemas.estimate_total import EstimateTotalRequest

logger = logging.getLogger(__name__)

TEMPLATE_PATH = os.path.join(os.path.dirname(__file__), "..", "..", "estimate_template.docx")

class EstimateDocxGenerator:
    def __init__(self, template_path=TEMPLATE_PATH):
        self.template_path = template_path
        self.font_name = "Open Sans"
        self.primary_color = 0x612d4b  # Wine color from HTML
        self.text_color = 0x333333     # Main text color
        self.desc_color = 0x555555     # Description color

    def _set_run_font(self, run, size_pt=Pt(10), bold=None, italic=None, color_rgb=None, underline=None):
        """Helper to consistently set font properties in a run."""
        rPr = run._element.get_or_add_rPr()
        
        # Set font name in both high-level and XML level
        run.font.name = self.font_name
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:ascii'), self.font_name)
        rFonts.set(qn('w:hAnsi'), self.font_name)
        rFonts.set(qn('w:cs'), self.font_name)

        if size_pt is not None:
            # size_pt is already in Pt/EMU if passed from Pt()
            # Directly assign to avoid double conversion bug
            run.font.size = size_pt
        if bold is not None:
            run.bold = bold
        if italic is not None:
            run.italic = italic
        if underline is not None:
            run.underline = underline
        if color_rgb is not None:
            run.font.color.rgb = RGBColor((color_rgb >> 16) & 0xff, (color_rgb >> 8) & 0xff, color_rgb & 0xff)

    def _format_currency(self, val):
        if val is None:
            return ""
        
        if isinstance(val, (int, float)):
            # US Format: 1,234.56
            s = f"{abs(val):,.2f}"
            if val < 0:
                return f"-$ {s}"
            return f"$ {s}"

        s = str(val).strip()
        if not s:
            return ""
        
        if not any(c.isdigit() for c in s):
            return s

        # Parse robustly to handle any input locale style
        parsed = self._parse_price(s)
        
        # Standardize on US Format: 1,234.56
        s_formatted = f"{abs(parsed):,.2f}"
        if parsed < 0:
            return f"-$ {s_formatted}"
        return f"$ {s_formatted}"

    def _parse_price(self, val):
        if not val:
            return 0.0
        if isinstance(val, (int, float)):
            return float(val)
        # Clean currency symbols and spaces
        clean = str(val).replace("$", "").replace(" ", "").strip()
        
        # Priority to US Format: 1,234.56
        # If there are both , and .
        if "," in clean and "." in clean:
            if clean.rfind(".") > clean.rfind(","): # US style: 1,234.56
                clean = clean.replace(",", "")
            else: # European style fallback: 1.234,56
                clean = clean.replace(".", "").replace(",", ".")
        elif "," in clean:
            # Ambiguous: 1,234 (thousands) or 1,23 (decimal)
            parts = clean.split(",")
            if len(parts[-1]) == 3: # Likely thousands: 1,000
                clean = clean.replace(",", "")
            else: # Likely decimal: 1,23
                clean = clean.replace(",", ".")
        
        try:
            return float(clean)
        except (ValueError, TypeError):
            return 0.0

    def _parse_percentage(self, val):
        if not val:
            return 0.0
        if isinstance(val, (int, float)):
            return float(val) / 100.0
        
        # Clean percentage symbol and spaces
        clean = str(val).replace("%", "").replace(" ", "").strip()
        
        # For percentages, if there's a comma, it's ALMOST ALWAYS a decimal separator 
        # (e.g., 6,350% in Connecticut or 20,00% Service Charge)
        if "," in clean and "." not in clean:
            clean = clean.replace(",", ".")
            
        # Use _parse_price logic for final conversion
        num = self._parse_price(clean)
        return num / 100.0

    def _parse_date_header(self, val):
        if not val:
            return datetime.min
        clean = str(val).strip()
        
        # Common formats used in AppSheet / English locale dates
        # e.g., "June, Wednesday 17 2026" -> "%B, %A %d %Y"
        formats = [
            "%B, %A %d %Y", # June, Wednesday 17 2026
            "%A, %B %d %Y", # Wednesday, June 17 2026
            "%B %d %Y",
            "%m/%d/%Y",
            "%d/%m/%Y",
            "%Y-%m-%d"
        ]
        for fmt in formats:
            try:
                return datetime.strptime(clean, fmt)
            except (ValueError, TypeError):
                continue
        
        return datetime.min

    def _replace_placeholders(self, doc, request: EstimateTotalRequest):
        """Replaces {{TAG}} placeholders in the document paragraphs and tables."""
        replacements = {
            "{{EVENT_NAME}}": request.event.name,
            "{{CLIENT_NAME}}": request.client.name,
            "{{CLIENT_ADDRESS}}": request.client.address,
            "{{CLIENT_EMAIL}}": request.client.email,
            "{{REPRESENTATIVE_NAME}}": request.client_representative.name,
            "{{REPRESENTATIVE_EMAIL}}": request.client_representative.email,
            "{{REPRESENTATIVE_PHONE}}": request.client_representative.formatted_phone,
            "{{EVENT_ADDRESS}}": request.event.address,
            "{{EVENT_CODE}}": request.event.code,
            "{{EVENT_START}}": request.event.date_formatted,
            "{{EVENT_END}}": request.event.end_date_formatted,
            "{{EVENT_GUESTS}}": str(request.event.guests),
            "{{SERVICE_CHARGE_RATE}}": request.financials.service_charge_rate,
            "{{TAX_NAME}}": request.financials.tax_name,
            "{{TAX_RATE}}": request.financials.tax_rate,
        }

        def process_paragraphs(paragraphs):
            for paragraph in paragraphs:
                for key, value in replacements.items():
                    if key in paragraph.text:
                        # 1. Try simple replacement in runs
                        found_in_run = False
                        for run in paragraph.runs:
                            if key in run.text:
                                run.text = run.text.replace(key, str(value or ""))
                                found_in_run = True
                        
                        # 2. If tag is split across runs, heal it by merging runs
                        if not found_in_run and len(paragraph.runs) > 1:
                            full_text = "".join(r.text for r in paragraph.runs)
                            if key in full_text:
                                paragraph.runs[0].text = full_text.replace(key, str(value or ""))
                                for i in range(1, len(paragraph.runs)):
                                    paragraph.runs[i].text = ""

        # Process main body paragraphs
        process_paragraphs(doc.paragraphs)

        # Process all tables (including nested ones)
        def process_tables(tables):
            for table in tables:
                for row in table.rows:
                    for cell in row.cells:
                        process_paragraphs(cell.paragraphs)
                        if cell.tables:
                            process_tables(cell.tables)

        process_tables(doc.tables)

        # Process Headers and Footers (where your left-side info is located)
        for section in doc.sections:
            for header in [section.header, section.first_page_header, section.even_page_header]:
                if header:
                    process_paragraphs(header.paragraphs)
                    process_tables(header.tables)
            for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
                if footer:
                    process_paragraphs(footer.paragraphs)
                    process_tables(footer.tables)
            
            # 3. Fallback: process raw XML for elements not caught in paragraphs (like floating Text Boxes)
            for part in [section.header, section.first_page_header, section.even_page_header, section.footer, section.first_page_footer, section.even_page_footer]:
                if part:
                    for t in part._element.iter("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t"):
                        if t.text:
                            for key, value in replacements.items():
                                if key in t.text:
                                    t.text = t.text.replace(key, str(value or ""))

    def generate_docx(self, request: EstimateTotalRequest) -> BytesIO:
        if not os.path.exists(self.template_path):
            raise FileNotFoundError(f"Template not found at {self.template_path}")

        doc = Document(self.template_path)
        self._replace_placeholders(doc, request)

        # De-duplicate meals to avoid repetitions (common in some data sources)
        unique_meals_map = {} # sig -> meal
        for m in request.meals:
            m_dict = m.model_dump()
            m_dict.pop('show_date_header', None)
            # Signature based on content to identify duplicates
            sig = str(sorted(m_dict.items(), key=lambda x: x[0]))
            if sig not in unique_meals_map:
                unique_meals_map[sig] = m
            else:
                # If this instance has show_date_header=True, prefer it
                if m.show_date_header:
                    unique_meals_map[sig] = m
        
        # Preserve original order as much as possible
        unique_meals = []
        seen_sigs = set()
        for m in request.meals:
            m_dict = m.model_dump()
            m_dict.pop('show_date_header', None)
            sig = str(sorted(m_dict.items(), key=lambda x: x[0]))
            if sig not in seen_sigs:
                seen_sigs.add(sig)
                unique_meals.append(unique_meals_map[sig])
        # Sort meals chronologically (stable sort keeps original sub-order within a day)
        unique_meals.sort(key=lambda m: self._parse_date_header(m.date_header))
        
        # Use unique_meals instead of request.meals for content generation
        # request.meals = unique_meals # We can also just use unique_meals below

        marker_para = None
        for p in doc.paragraphs:
            if "[DYNAMIC_CONTENT_START]" in p.text:
                marker_para = p
                break

        if not marker_para:
            logger.warning("Marker [DYNAMIC_CONTENT_START] not found. Appending to end.")

        def add_p(text="", alignment=None, space_after=Pt(0), space_before=Pt(0), bold=False, italic=False, size=Pt(10), color=0x333333, underline=False):
            if marker_para:
                p = marker_para.insert_paragraph_before(text)
            else:
                p = doc.add_paragraph(text)
            
            if alignment: p.alignment = alignment
            p.paragraph_format.space_after = space_after
            p.paragraph_format.space_before = space_before
            
            if p.runs:
                # Use current run if it exists, otherwise add one
                run = p.runs[0]
            else:
                run = p.add_run(text)
            
            # The add_run(text) might result in double text if p.runs already had text
            # but usually add_p(text) with marker_para doesn't have runs yet.
            if not p.runs and text:
                run = p.add_run(text)

            self._set_run_font(run, size_pt=size, bold=bold, italic=italic, color_rgb=color, underline=underline)
            
            return p

        def add_hr():
            p = add_p(space_after=Pt(3), space_before=Pt(0), size=Pt(1))
            p.paragraph_format.line_spacing = Pt(1)
            p_pr = p._element.get_or_add_pPr()
            # Check if pBdr already exists to prevent duplication
            p_bdr = p_pr.find(qn('w:pBdr'))
            if p_bdr is None:
                p_bdr = OxmlElement('w:pBdr')
                p_pr.insert(0, p_bdr)
            
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), '000000')
            p_bdr.append(bottom)

        # --- EVENT HEADER ---
        if request.event.name:
            add_p(request.event.name, alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=Pt(12), color=self.primary_color, space_after=Pt(4))
        if request.event.address:
            add_p(request.event.address, alignment=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(9), color=self.text_color, space_after=Pt(12))

        # --- MENU SECTION ---
        add_p("MENUS", bold=True, size=Pt(10), color=self.primary_color, space_after=Pt(0))
        add_p(request.event.date_formatted, space_after=Pt(0))

        if request.event.dietary_restrictions:
            add_p("Dietary Restrictions", bold=True, size=Pt(10), color=self.primary_color, space_after=Pt(0))
            add_p(request.event.dietary_restrictions, space_after=Pt(0))

        printed_dates_menu = set()
        for meal in unique_meals:
            norm_date = (meal.date_header or "").strip()
            if norm_date not in printed_dates_menu:
                printed_dates_menu.add(norm_date)
                add_p(meal.date_header, bold=True, space_before=Pt(6))
                add_hr()
            
            p = add_p("", space_before=Pt(8))
            if meal.time_range:
                r_cat = p.add_run(meal.category_name.upper() + ": ")
                self._set_run_font(r_cat, bold=True, size_pt=Pt(10), color_rgb=self.primary_color)
                r_time = p.add_run(meal.time_range)
                self._set_run_font(r_time, bold=False, size_pt=Pt(10), color_rgb=self.text_color)
            else:
                r_cat = p.add_run(meal.category_name.upper())
                self._set_run_font(r_cat, bold=True, size_pt=Pt(10), color_rgb=self.primary_color)
            
            if meal.provide_by_client:
                p = add_p("◽ Provided by client", space_before=Pt(4))
                p.paragraph_format.left_indent = Cm(0.5)
                continue

            if meal.description:
                add_p(meal.description, italic=True, space_before=Pt(0))

            # --- MEAL LEVEL ---
            # Group subcategories by name to avoid repetitions (e.g. multiple "Including" or "Vegan" blocks)
            grouped_subs = [] 
            for i in range(1, 13):
                s_name = getattr(meal, f"subcategory_{i}_name", "").strip()
                s_desc = getattr(meal, f"subcategory_{i}_description", "").strip()
                s_items = getattr(meal, f"subcategory_{i}_items", [])

                if not s_name and not s_items:
                    continue

                found = False
                if s_name:
                    for gs in grouped_subs:
                        if gs['name'] == s_name:
                            # Append unique items
                            existing_names = {it.name for it in gs['items']}
                            for it in s_items:
                                if it.name not in existing_names:
                                    gs['items'].append(it)
                                    existing_names.add(it.name)
                            if s_desc and s_desc not in gs['desc']:
                                gs['desc'] = (gs['desc'] + " " + s_desc).strip()
                            found = True
                            break
                
                if not found:
                    grouped_subs.append({
                        'name': s_name,
                        'desc': s_desc,
                        'items': list(s_items)
                    })

            for gs in grouped_subs:
                sub_name = gs['name']
                sub_desc = gs['desc']
                sub_items = gs['items']

                if sub_name:
                    sub_p = add_p(sub_name, bold=True, space_before=Pt(6))
                    sub_p.runs[0].underline = True

                if sub_desc:
                    add_p(sub_desc, size=Pt(10), italic=True)

                for item in sub_items:
                    # Construct the main item line
                    item_p = add_p(space_after=Pt(2))
                    item_p.paragraph_format.left_indent = Cm(0.8)
                    item_p.paragraph_format.first_line_indent = Cm(-0.4)
                    
                    r_bullet = item_p.add_run("◽ ")
                    self._set_run_font(r_bullet, bold=True)
                    
                    r_name = item_p.add_run(item.name)
                    self._set_run_font(r_name, bold=False, underline=False)

                    if item.description:
                        desc_p = add_p(item.description, size=Pt(10), italic=True, color=self.desc_color, space_after=Pt(4))
                        desc_p.paragraph_format.left_indent = Cm(1.2)

        # --- FINANCIAL SECTION ---
        # Force a page break before financials if needed, or just a big spacer
        add_p(space_before=Pt(10))
        add_p("PROPOSAL OF SERVICES", bold=True, size=Pt(10), color=self.primary_color, space_after=Pt(0), space_before=Pt(10))
        add_p(request.event.end_date_formatted, space_after=Pt(0)) # HTML uses End Event here
        #add_p(size=Pt(8))

        # 1. Food Service
        add_p("Food Service", bold=True, size=Pt(10), color=self.primary_color, space_after=Pt(0), space_before=Pt(10))
        add_p(f"Based on {request.event.guests} Guests", size=Pt(10), italic=True, space_before=Pt(0))
        
        # Pre-calculate daily food totals to ensure accuracy (ignoring "Provided by client" items)
        daily_food_totals = {}
        for m in unique_meals:
            if not m.provide_by_client:
                val = self._parse_price(m.total_category_precio)
                daily_food_totals[m.date_header] = daily_food_totals.get(m.date_header, 0.0) + val

        meals_by_date = {}
        for m in unique_meals:
            date_key = (m.date_header or "").strip()
            if date_key not in meals_by_date:
                meals_by_date[date_key] = []
            meals_by_date[date_key].append(m)

        for date_header, day_meals in meals_by_date.items():
            orig_date_header = day_meals[0].date_header
            add_p(orig_date_header, bold=True, space_before=Pt(6), space_after=Pt(0))
            add_hr()
            
            for meal in day_meals:
                p = add_p(space_after=Pt(2))
                p.paragraph_format.tab_stops.add_tab_stop(Cm(16.5), WD_TAB_ALIGNMENT.RIGHT)
                r_label = p.add_run(re.sub(r'\s+', ' ', (meal.category_precio_guest or "")).strip())
                self._set_run_font(r_label)
                if not meal.provide_by_client:
                    val = self._parse_price(meal.total_category_precio)
                    if abs(val) >= 0.01:
                        r_spacer = p.add_run("\t")
                        self._set_run_font(r_spacer)
                        r_val = p.add_run(self._format_currency(val))
                        self._set_run_font(r_val, bold=False)
                else:
                    r_spacer = p.add_run("\t")
                    self._set_run_font(r_spacer)
                    r_client = p.add_run("Provided by client")
                    self._set_run_font(r_client, italic=True)

            # Place daily total at the bottom in bold, aligned right
            total_val = daily_food_totals.get(orig_date_header, 0.0)
            if total_val >= 0:
                p_tot = add_p(space_after=Pt(4))
                p_tot.paragraph_format.tab_stops.add_tab_stop(Cm(16.5), WD_TAB_ALIGNMENT.RIGHT)
                r_spacer = p_tot.add_run("\t")
                self._set_run_font(r_spacer)
                r_tot = p_tot.add_run(self._format_currency(total_val))
                self._set_run_font(r_tot, bold=True)

        # 2. Labor
        if request.labor_services:
            add_p("Labor Service Fees", bold=True, size=Pt(10), color=self.primary_color, space_before=Pt(15), space_after=Pt(0))
            
            # 1. De-duplicate labor services first to avoid repetitions (with normalization)
            seen_labor = set()
            unique_labor = []
            for labor in request.labor_services:
                norm_date = (labor.date_header or "").strip()
                norm_hours = self._parse_price(labor.hours)
                norm_name = (labor.name or "").strip()
                norm_total = self._parse_price(labor.total)
                
                key = (norm_date, norm_hours, norm_name, norm_total)
                if key not in seen_labor:
                    seen_labor.add(key)
                    unique_labor.append(labor)

            # 2. Group labor by date and hours (using normalized comparison)
            labor_groups = []
            for labor in unique_labor:
                found = False
                norm_labor_date = (labor.date_header or "").strip()
                norm_labor_hours = self._parse_price(labor.hours)
                
                for g in labor_groups:
                    norm_g_date = (g['date'] or "").strip()
                    norm_g_hours = self._parse_price(g['hours'])
                    
                    if norm_g_date == norm_labor_date and norm_g_hours == norm_labor_hours:
                        g['items'].append(labor)
                        found = True
                        break
                if not found:
                    labor_groups.append({
                        'date': labor.date_header,
                        'hours': labor.hours,
                        'show_date': labor.show_date_header,
                        'items': [labor]
                    })

            # Sort labor groups chronologically
            labor_groups.sort(key=lambda g: self._parse_date_header(g['date']))

            printed_dates_labor = set()
            for group in labor_groups:
                norm_date = (group['date'] or "").strip()
                if norm_date not in printed_dates_labor:
                    printed_dates_labor.add(norm_date)
                    add_p(group['date'], bold=True, space_before=Pt(6), space_after=Pt(0))
                    add_hr()
                
                # Paragraph for the description (continuous text)
                p_desc = add_p(space_after=Pt(0))
                
                # 1. Header (Italic with Bold Hours)
                r_header_prefix = p_desc.add_run("Staff suggested based on ")
                self._set_run_font(r_header_prefix, italic=True)
                
                r_hours = p_desc.add_run(f"{group['hours']}")
                self._set_run_font(r_hours, italic=True, bold=True)
                
                r_header_suffix = p_desc.add_run(" hours of labor. ")
                self._set_run_font(r_header_suffix, italic=True)
                
                # 2. Names concatenated (Normal weight)
                names_str = ", ".join([re.sub(r'\s+', ' ', (item.name or "")).strip() for item in group['items']])
                r_names = p_desc.add_run(names_str)
                self._set_run_font(r_names, bold=False)
                
                # 3. Total on a new line (aligned right)
                total_val = sum(self._parse_price(item.total) for item in group['items'])
                p_total = add_p(self._format_currency(total_val), alignment=WD_ALIGN_PARAGRAPH.RIGHT, bold=True, space_after=Pt(4))
                # Explicitly set color for the total
                self._set_run_font(p_total.runs[0], bold=True, color_rgb=0x000000)

        # 3. Extras
        if request.extras_events:
            add_p("Extras Services", bold=True, size=Pt(10), color=self.primary_color, space_before=Pt(15), space_after=Pt(0))
            
            # De-duplicate extras events (preserve order, with normalization)
            seen_extras = set()
            unique_extras = []
            for extra in request.extras_events:
                norm_date = (extra.date_header or "").strip()
                norm_name = (extra.name or "").strip()
                norm_rental = (extra.name_rental or "").strip()
                norm_sales = (extra.name_sales or "").strip()
                norm_total = self._parse_price(extra.total)
                
                key = (norm_date, extra.is_rental, extra.is_sales, norm_name, norm_rental, norm_sales, norm_total, extra.provide_by_client)
                if key not in seen_extras:
                    seen_extras.add(key)
                    unique_extras.append(extra)

            # Sort extras chronologically
            unique_extras.sort(key=lambda ex: self._parse_date_header(ex.date_header))

            printed_dates_extras = set()
            for extra in unique_extras:
                norm_date = (extra.date_header or "").strip()
                if norm_date not in printed_dates_extras:
                    printed_dates_extras.add(norm_date)
                    add_p(extra.date_header, bold=True, space_after=Pt(0))
                    add_hr()
                
                if extra.is_rental:
                    add_p("Rentals", bold=True, space_before=Pt(6))
                
                if extra.is_sales:
                    add_p("Sales", bold=True, space_before=Pt(6))

                p = add_p(space_after=Pt(2))
                p.paragraph_format.tab_stops.add_tab_stop(Cm(16.5), WD_TAB_ALIGNMENT.RIGHT)
                # Determine name based on flags
                display_name = extra.name
                if not extra.provide_by_client:
                    if extra.is_rental: display_name = extra.name_rental
                    elif extra.is_sales: display_name = extra.name_sales

                if extra.provide_by_client:
                    txt = f"{display_name}\tProvide by the client"
                else:
                    txt = f"{display_name}\t{self._format_currency(extra.total)}"
                p_extra = p.add_run(txt)
                self._set_run_font(p_extra, bold=True)

        # 4. Final Summary
        add_p("Cost of Balance", bold=True, size=Pt(10), color=self.primary_color, space_before=Pt(10))
        fin = request.financials
        # --- RE-CALCULATION ENGINE (Excel Model) ---
        # 1. Base Components
        real_food_total = sum(daily_food_totals.values())
        
        real_labor_total = 0.0
        if request.labor_services:
            for group in labor_groups:
                real_labor_total += sum(self._parse_price(item.total) for item in group['items'])
            
        real_extras_sales_total = 0.0
        real_extras_rentals_total = 0.0
        if request.extras_events:
            for ex in unique_extras:
                if not ex.provide_by_client:
                    val = self._parse_price(ex.total)
                    if ex.is_sales: real_extras_sales_total += val
                    if ex.is_rental: real_extras_rentals_total += val

        real_gratuity = self._parse_price(fin.gratuity)
        real_discount = self._parse_price(fin.discount)
        real_donation = self._parse_price(fin.donation)

        # 2. SubTotal 1 & Taxes
        # SubTotal 1 = Food + Labor + Extras Sales + Gratuity - Discount - Donation
        subtotal_1 = (real_food_total + real_labor_total + real_extras_sales_total + 
                      real_gratuity - abs(real_discount) - abs(real_donation))
        
        tax_rate = self._parse_percentage(fin.tax_rate)
        real_tax = subtotal_1 * tax_rate
        subtotal_2 = subtotal_1 + real_tax

        # 3. Service Charge & SubTotal 4
        # Service Charge = (Food + Labor) * Rate
        service_charge_rate = self._parse_percentage(fin.service_charge_rate)
        real_service_charge = (real_food_total + real_labor_total) * service_charge_rate
        
        # Subtotal 4 = Subtotal 2 + Extras Rentals + Service Charge
        subtotal_4 = subtotal_2 + real_extras_rentals_total + real_service_charge

        # 4. Credit Card & Final Total
        cc_rate = self._parse_percentage(fin.credit_card_percent)
        real_cc_fee = subtotal_4 * cc_rate
        real_grand_total = subtotal_4 + real_cc_fee

        # --- RENDER SUMMARY ---
        summary_items = [
            ("Food", real_food_total, True),
            ("Labor Cost", real_labor_total, True),
            ("Extras Services (Sales)", real_extras_sales_total, True),
            ("Gratuity", real_gratuity, False),
            ("Discount", -abs(real_discount), False),
            ("Donation", -abs(real_donation), False),
            # ("Event Subtotal (Pre-Tax)", subtotal_1, True, True), # Removed from view per request
            (f"{fin.tax_rate} {fin.tax_name}", real_tax, True),
            ("Subtotal after Taxes", subtotal_2, True, True),
            ("Extras Services (Rentals)", real_extras_rentals_total, True),
            (f"{fin.service_charge_rate} Service Charge", real_service_charge, True),
            ("Total Estimated Amount", subtotal_4, True, True),
            ("Credit Card Fee", real_cc_fee, False),
        ]

        for item in summary_items:
            label = item[0]
            val = item[1]
            is_bold_item = item[3] if len(item) > 3 else False

            # Skip all items with value 0.00 except for critical grand totals
            is_core_total = label in ["Subtotal after Taxes", "Total Estimated Amount"]
            if not is_core_total and abs(val) < 0.01:
                continue
                
            p = add_p(space_after=Pt(2))
            p.paragraph_format.tab_stops.add_tab_stop(Cm(16.5), WD_TAB_ALIGNMENT.RIGHT)
            
            r_label = p.add_run(label)
            self._set_run_font(r_label, bold=is_bold_item)
            
            r_tab = p.add_run("\t")
            self._set_run_font(r_tab)
            
            formatted_val = self._format_currency(val)
            r_val = p.add_run(formatted_val)
            self._set_run_font(r_val, bold=is_bold_item)

        # Final Total Line
        total_p = add_p(space_after=Pt(2), space_before=Pt(8))
        total_p.paragraph_format.tab_stops.add_tab_stop(Cm(16.5), WD_TAB_ALIGNMENT.RIGHT)
        
        r_total = total_p.add_run(f"Final Balance Due\t{self._format_currency(real_grand_total)}")
        self._set_run_font(r_total, bold=True, size_pt=Pt(10), color_rgb=self.primary_color)
        
        p_pr = total_p._element.get_or_add_pPr()
        p_bdr = p_pr.find(qn('w:pBdr'))
        if p_bdr is None:
            p_bdr = OxmlElement('w:pBdr')
            p_pr.insert(0, p_bdr)
        
        top = OxmlElement('w:top')
        top.set(qn('w:val'), 'single')
        top.set(qn('w:sz'), '12')
        top.set(qn('w:color'), '612D4B')
        p_bdr.append(top)

        if marker_para:
            p_element = marker_para._element
            p_element.getparent().remove(p_element)

        docx_stream = BytesIO()
        doc.save(docx_stream)
        docx_stream.seek(0)
        return docx_stream
