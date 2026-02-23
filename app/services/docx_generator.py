import os
from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from app.schemas.menu import MenuRequest

TEMPLATE_PATH = os.path.join(os.path.dirname(__file__), "..", "..", "template.docx")

def get_font_sizes(item_count: int):
    # Dynamic styling based on item count, matching the logic in the provided HTML CSS
    if item_count <= 3:
        return {"subcat": 14, "menu": 12}
    elif item_count <= 6:
        return {"subcat": 12, "menu": 11}
    elif item_count <= 9:
        return {"subcat": 11, "menu": 10}
    else:  # 10+
        return {"subcat": 10, "menu": 9}

def format_paragraph(paragraph, text: str, font_name: str, size: int, bold: bool = False, color: RGBColor = None):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    return run

def generate_menu_docx(request: MenuRequest) -> BytesIO:
    # Use the user-provided template if it exists, otherwise create a blank one
    if os.path.exists(TEMPLATE_PATH):
        doc = Document(TEMPLATE_PATH)
    else:
        doc = Document()
        
    # Center content vertically for all sections
    for section in doc.sections:
        # Use low-level XML to ensure compatibility across python-docx versions
        section._sectPr.vAlign_val = 'center'
        
    meals_count = len(request.all_meals)
    
    for idx, meal in enumerate(request.all_meals):
        # Header section
        title_para = doc.add_paragraph()
        format_paragraph(title_para, meal.categoria, "Georgia", 24, bold=True, color=RGBColor(0x5a, 0x2d, 0x5a)) # #5a2d5a
        
        date_para = doc.add_paragraph()
        format_paragraph(date_para, meal.fecha, "Georgia", 12, color=RGBColor(0x33, 0x33, 0x33))
        
        desc_para = doc.add_paragraph()
        format_paragraph(desc_para, meal.descripcion, "Georgia", 13, color=RGBColor(0x33, 0x33, 0x33))
        
        # Add some spacing before items
        doc.add_paragraph()
        
        # Filter out items where subcat is empty
        valid_items = [item for item in meal.items if item.subcat.strip()]
        
        sizes = get_font_sizes(len(valid_items))
        
        for item in valid_items:
            subcat_para = doc.add_paragraph()
            run = format_paragraph(subcat_para, item.subcat, "Georgia", sizes["subcat"], bold=True, color=RGBColor(0x00, 0x00, 0x00))
            run.underline = True 
            
            if item.menu.strip():
                menu_para = doc.add_paragraph()
                format_paragraph(menu_para, item.menu, "Georgia", sizes["menu"], color=RGBColor(0x33, 0x33, 0x33))
        
        # Add page break if it's not the last meal
        if idx < meals_count - 1:
            doc.add_page_break()
            
    # Save to memory stream
    docx_stream = BytesIO()
    doc.save(docx_stream)
    docx_stream.seek(0)
    
    return docx_stream
