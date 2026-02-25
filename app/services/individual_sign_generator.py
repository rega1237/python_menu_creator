import os
import logging
from io import BytesIO
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from app.schemas.individual_menu import IndividualSignRequest

logger = logging.getLogger(__name__)

LOGO_PATH = os.path.join(os.path.dirname(__file__), "..", "static", "logo_fdce.png")

class ItemData:
    def __init__(self, name, description, diet_options):
        self.name = name
        self.description = description
        self.diet_options = diet_options


def format_cell(cell, item):
    """Formats a table cell as a label."""
    # Set cell dimensions (approx 8cm x 5cm)
    # Note: python-docx doesn't always strictly enforce cell width if table is auto-layout
    # but we will try.
    
    # Clear cell
    for paragraph in cell.paragraphs:
        p = paragraph._element
        p.getparent().remove(p)
        
    # Vertical center
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    # Name (Destacado)
    name_para = cell.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = name_para.add_run(item.name.upper())
    run.font.name = "Arial"
    run.font.size = Pt(22)
    run.bold = True
    run.font.color.rgb = RGBColor(0x5a, 0x2d, 0x5a)
    
    # Description
    desc_para = cell.add_paragraph()
    desc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = desc_para.add_run(item.description)
    run.font.name = "Arial"
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    
    # Diet Options
    if item.diet_options and item.diet_options.strip() and item.diet_options.strip() != '""' and item.diet_options.strip() != '"".""':
        diet_para = cell.add_paragraph()
        diet_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = diet_para.add_run(f"({item.diet_options.strip().upper()})")
        run.font.name = "Arial"
        run.font.size = Pt(12)
        run.bold = True
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Logo (Bottom right)
    if os.path.exists(LOGO_PATH):
        logo_para = cell.add_paragraph()
        logo_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = logo_para.add_run()
        run.add_picture(LOGO_PATH, width=Cm(1.2))

def create_grid_page(doc):
    """Creates a 3x2 table grid for a page."""
    table = doc.add_table(rows=3, cols=2)
    # table.style = 'Table Grid' # Change to None or custom style for no borders if requested
    # We want borders to match the user's design
    
    # Set heights to approx 5cm to 6cm
    for row in table.rows:
        row.height = Cm(5.5)
        
    # Set widths to 8cm
    for col in table.columns:
        col.width = Cm(9) # Slightly wider than 8cm to fill page
        
    # Apply borders to all cells
    borders_xml = """
        <w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
            <w:top w:val="single" w:sz="12" w:space="0" w:color="5a2d5a"/>
            <w:left w:val="single" w:sz="12" w:space="0" w:color="5a2d5a"/>
            <w:bottom w:val="single" w:sz="12" w:space="0" w:color="5a2d5a"/>
            <w:right w:val="single" w:sz="12" w:space="0" w:color="5a2d5a"/>
        </w:tcBorders>
    """
    for row in table.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            
            top = OxmlElement('w:top')
            top.set(qn('w:val'), 'single')
            top.set(qn('w:sz'), '12')
            top.set(qn('w:space'), '0')
            top.set(qn('w:color'), '5a2d5a')
            
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '12')
            bottom.set(qn('w:space'), '0')
            bottom.set(qn('w:color'), '5a2d5a')
            
            left = OxmlElement('w:left')
            left.set(qn('w:val'), 'single')
            left.set(qn('w:sz'), '12')
            left.set(qn('w:space'), '0')
            left.set(qn('w:color'), '5a2d5a')
            
            right = OxmlElement('w:right')
            right.set(qn('w:val'), 'single')
            right.set(qn('w:sz'), '12')
            right.set(qn('w:space'), '0')
            right.set(qn('w:color'), '5a2d5a')
            
            tcBorders.append(top)
            tcBorders.append(bottom)
            tcBorders.append(left)
            tcBorders.append(right)
            tcPr.append(tcBorders)
        
    # Merge cells in the middle row for the center element
    table.cell(1, 0).merge(table.cell(1, 1))
    
    return table

def generate_individual_signs_docx(request: IndividualSignRequest) -> BytesIO:
    doc = Document()
    
    # Set Page Orientation to Landscape
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    # Update height/width for landscape
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height
    
    # Margins (1cm)
    section.top_margin = Cm(1)
    section.bottom_margin = Cm(1)
    section.left_margin = Cm(1)
    section.right_margin = Cm(1)
    
    items = []
    
    def process_and_add_item(name, desc, diet):
        if name and name.strip():
            # Check if name contains commas and split it
            if "," in name:
                name_parts = [p.strip() for p in name.split(",") if p.strip()]
                # If desc has commas, split it too, to map 1:1 with names. 
                # Otherwise, treat it as empty or use the whole thing if it's only 1 part (which shouldn't happen if names > 1, but just in case)
                desc_parts = []
                if desc:
                    desc_parts = [p.strip() for p in desc.split(",")]
                
                diet_parts = []
                if diet:
                    # Sometimes diet is "(GF), (VG)" or just "GF, VG". We split by comma if present.
                    # But diet often applies to all, or is mapped 1:1. We assume 1:1 if it has commas.
                    diet_parts = [p.strip() for p in diet.split(",")]

                for i, part in enumerate(name_parts):
                    # Try to get the matching description, fallback to empty string
                    part_desc = desc_parts[i] if i < len(desc_parts) else ""
                    # If there's only one description but many names, it might be a global description. 
                    # But per user request, we want to split. If it's missing, leave it blank.
                    if not part_desc and len(desc_parts) == 1:
                         # Edge case: If there's exactly 1 description but 4 names, do we repeat or leave blank?
                         # The user specifically complained it repeated. So we leave it blank if no 1:1 match.
                         # Actually, wait. If `Dried Apricots, Dried Cramberri` is the desc, it splits to 2.
                         # If `APRICOTS, CRAMBERRI` is the name, it splits to 2. Perfect match.
                         part_desc = desc_parts[0] if len(desc_parts) == 1 else ""

                    part_diet = diet_parts[i] if i < len(diet_parts) else ""
                    if not part_diet and len(diet_parts) == 1:
                        part_diet = diet_parts[0]
                        
                    items.append(ItemData(
                        part,
                        part_desc, 
                        part_diet
                    ))
            else:
                items.append(ItemData(
                    name.strip(), 
                    desc.strip() if desc else "", 
                    diet.strip() if diet else ""
                ))

    for meal in request.meals:
        process_and_add_item(meal.menu_name, meal.menu_desc, meal.menu_diet)
            
        for i in range(1, 11):
            name = getattr(meal, f"menu_{i}_name")
            desc = getattr(meal, f"menu_{i}_desc")
            diet = getattr(meal, f"menu_{i}_diet")
            process_and_add_item(name, desc, diet)
    
    total_items = len(items)
    
    for i in range(0, total_items, 5):
        if i > 0:
            doc.add_page_break()
            
        table = create_grid_page(doc)
        page_items = items[i:i+5]
        
        # Pattern:
        # R0C0 (Item 1), R0C1 (Item 2)
        # R1 (Merged) (Item 3)
        # R2C0 (Item 4), R2C1 (Item 5)
        
        # Mapping indices to table cells
        # cell_map: (page_index) -> (row, col)
        cell_map = {
            0: (0, 0),
            1: (0, 1),
            2: (1, 0), # This is the merged center cell
            3: (2, 0),
            4: (2, 1)
        }
        
        for idx, item in enumerate(page_items):
            row, col = cell_map[idx]
            cell = table.cell(row, col)
            format_cell(cell, item)
            
    # Save to memory stream
    docx_stream = BytesIO()
    doc.save(docx_stream)
    docx_stream.seek(0)
    
    return docx_stream
